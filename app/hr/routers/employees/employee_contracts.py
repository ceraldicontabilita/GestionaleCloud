"""
Employee Contracts Router - Gestione contratti dipendenti.
"""
from fastapi import APIRouter, HTTPException, Body, UploadFile, File, Depends
from fastapi.responses import FileResponse, Response, StreamingResponse
from typing import Dict, Any, List, Optional
from datetime import datetime, timezone
import logging
import os
import io
import ssl
import smtplib
import base64
import uuid
import shutil
from email.message import EmailMessage
from docx import Document
import tempfile

from app.hr.database import Database, Collections
from app.hr.utils.error_handler import handle_errors
from app.hr.services.openapi_signature import (
    get_client, OpenAPIConfigError, OpenAPIError,
)
from app.hr.services.docx_converter import docx_to_pdf, DocxConversionError

logger = logging.getLogger(__name__)
router = APIRouter()

COLL_TEMPLATES = "contract_templates"   # template .docx persistenti (MongoDB-first)

# Directory effimere (solo file temporanei di lavorazione): /tmp è scrivibile su Render.
CONTRACTS_DIR = "/tmp/uploads/contracts"
TEMPLATES_DIR = "/tmp/uploads/contract_templates"

# Available contract types
CONTRACT_TYPES = [
    {"id": "determinato", "name": "Contratto a Tempo Determinato", "filename": "Contratto derminato.docx"},
    {"id": "indeterminato", "name": "Contratto a Tempo Indeterminato", "filename": "Contratto indetermionato.docx"},
    {"id": "part_time_det", "name": "Contratto Part-Time Determinato", "filename": "Contratto part_time determinato.docx"},
    {"id": "part_time_ind", "name": "Contratto Part-Time Indeterminato", "filename": "Contratto part_time indeterminato.docx"},
    {"id": "informativa_152", "name": "Informativa D.Lgs. 152/1997", "filename": "INFORMATIVA AI SENSI DEL D.LGS. 152-1997.docx"},
    {"id": "informativa_privacy", "name": "Informativa Privacy", "filename": "Informativa-Privacy.docx"},
    {"id": "regolamento", "name": "Regolamento Interno Aziendale", "filename": "REGOLAMENTO INTERNO AZIENDALE.docx"},
    {"id": "richiesta_ferie", "name": "Richiesta Ferie", "filename": "RICHIESTA FERIE.docx"},
    {"id": "riduzione_orario", "name": "Accordo Riduzione Orario / Solidarietà", "filename": "Accordo riduzione orario.docx"},
]


def ensure_dirs():
    """Create directories if they don't exist (best-effort)."""
    for d in (CONTRACTS_DIR, TEMPLATES_DIR):
        try:
            os.makedirs(d, exist_ok=True)
        except Exception:
            pass


async def _resolve_template(ct: Dict[str, str]) -> str:
    """Risolve il percorso del template .docx.

    Priorità MongoDB-first: se il template è salvato in `contract_templates`
    (persistente tra i deploy), lo scrive in un file temporaneo e ritorna quel
    path; altrimenti usa il file su disco in TEMPLATES_DIR (effimero su Render).
    """
    ensure_dirs()
    db = Database.get_db()
    doc = await db[COLL_TEMPLATES].find_one({"tipo": ct["id"]}, {"_id": 0, "file_data": 1})
    if doc and doc.get("file_data"):
        tmp = tempfile.mktemp(suffix=".docx")
        with open(tmp, "wb") as f:
            f.write(base64.b64decode(doc["file_data"]))
        return tmp
    disk = os.path.join(TEMPLATES_DIR, ct["filename"])
    if os.path.exists(disk):
        return disk
    raise HTTPException(404, f"Template non caricato: {ct['name']}. Caricalo dalla sezione Assunzione.")


def _to_float(val: Any) -> Optional[float]:
    """Converte in float un valore numerico tollerando la virgola decimale IT."""
    if val is None or val == "":
        return None
    try:
        return float(str(val).replace("€", "").replace(",", ".").strip())
    except (ValueError, TypeError):
        return None


def _fmt_euro(val: Optional[float]) -> str:
    """Formatta un importo in stile italiano (1.234,56) senza simbolo."""
    if val is None:
        return "______"
    return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def compute_stipendio_mensile(stipendio_orario: Any, ore_settimanali: Any) -> Optional[float]:
    """Calcola il lordo mensile teorico: paga oraria × ore settimanali × 52 / 12."""
    orario = _to_float(stipendio_orario)
    ore = _to_float(ore_settimanali)
    if orario is None or ore is None:
        return None
    return round(orario * ore * 52 / 12, 2)


def fill_contract_template(template_path: str, employee_data: Dict[str, Any]) -> str:
    """
    Fill contract template with employee data.
    Replaces specific text patterns with employee data.

    Supporta due meccanismi di segnaposto, combinabili nello stesso .docx:
      1. Puntini di sospensione (… ……) compilati per posizione (legacy).
      2. Segnaposto nominali `{{chiave}}` (es. {{ore_settimanali}},
         {{stipendio_mensile}}, {{periodo_prova}}) — il modo consigliato per i
         nuovi campi CCNL Turismo, indipendente dal layout del documento.
    """
    doc = Document(template_path)

    # Build full name
    nome_completo = employee_data.get("nome_completo", "")
    if not nome_completo:
        nome_completo = f"{employee_data.get('cognome', '')} {employee_data.get('nome', '')}".strip()

    # Campi CCNL Pubblici Esercizi / Turismo (H05Y) — parametrici, non inventati.
    ore_settimanali = employee_data.get("ore_settimanali") or "40"
    stipendio_orario = employee_data.get("stipendio_orario") or employee_data.get("salary")
    mensile = compute_stipendio_mensile(stipendio_orario, ore_settimanali)
    ferie_giorni = employee_data.get("ferie_giorni") or "26"
    periodo_prova = employee_data.get("periodo_prova") or ""
    ticket_attivo = bool(employee_data.get("ticket_buono"))
    ticket_importo = employee_data.get("ticket_importo")
    tredicesima = employee_data.get("tredicesima", True)
    quattordicesima = employee_data.get("quattordicesima", True)

    # Frasi pronte da inserire nel .docx tramite segnaposto nominali.
    if ticket_attivo:
        _imp = _to_float(ticket_importo)
        ticket_txt = (
            f"Buono pasto di euro {_fmt_euro(_imp)} giornalieri, riconosciuto dopo 1 anno di servizio."
            if _imp is not None else
            "Buono pasto giornaliero riconosciuto dopo 1 anno di servizio."
        )
    else:
        ticket_txt = "Non previsto."
    mensilita_lista = []
    if tredicesima:
        mensilita_lista.append("13ª (corrisposta a dicembre)")
    if quattordicesima:
        mensilita_lista.append("14ª (corrisposta a luglio)")
    mensilita_txt = " e ".join(mensilita_lista) if mensilita_lista else "12 mensilità"

    # All values to replace
    def g(*keys, default="______"):
        """Primo valore non vuoto tra le chiavi date; evita di scrivere 'None'."""
        for k in keys:
            v = employee_data.get(k)
            if v not in (None, "", "None"):
                return v
        return default

    def _fmt_date(v, default="______"):
        """Formatta una data in gg/mm/aaaa; tollera ISO e valori già formattati."""
        if v in (None, "", "None"):
            return default
        s = str(v)
        try:
            return datetime.fromisoformat(s.replace("Z", "")).strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            return s

    data_values = {
        "nome_completo": nome_completo or "______",
        "cognome": g("cognome"),
        "nome": g("nome"),
        "codice_fiscale": g("codice_fiscale", "cf"),
        "data_nascita": _fmt_date(g("data_nascita", default=None)),
        "luogo_nascita": g("luogo_nascita", "comune_nascita", "citta_nascita"),
        "indirizzo": g("indirizzo", "residenza"),
        "mansione": g("mansione", "qualifica"),
        "livello": g("livello"),
        "qualifica": g("qualifica", "mansione"),
        "stipendio_orario": str(stipendio_orario) if stipendio_orario not in (None, "") else "______",
        "data_inizio": _fmt_date(g("data_inizio", "hire_date", default=None)),
        "data_fine": _fmt_date(g("data_fine", default=None), default=""),
        # Nuovi campi CCNL Turismo (usabili come {{chiave}} nei .docx)
        "ore_settimanali": str(ore_settimanali),
        "stipendio_mensile": _fmt_euro(mensile),
        "ferie_giorni": str(ferie_giorni),
        "periodo_prova": str(periodo_prova) if periodo_prova not in (None, "") else "______",
        "ticket": ticket_txt,
        "mensilita": mensilita_txt,
    }

    # Decorrenza: senza data fine (indeterminato) niente "al ..." in coda.
    _df = data_values["data_fine"]
    _decorr = f"decorre dal {data_values['data_inizio']}" + (f" al {_df}" if _df else "")

    # Alias accettati per i segnaposto nominali (tolleranza sui nomi nel .docx).
    named_aliases = {
        "ore": "ore_settimanali",
        "ore_lavoro": "ore_settimanali",
        "mensile": "stipendio_mensile",
        "stipendio_mese": "stipendio_mensile",
        "paga_mensile": "stipendio_mensile",
        "paga_oraria": "stipendio_orario",
        "ferie": "ferie_giorni",
        "prova": "periodo_prova",
        "buono_pasto": "ticket",
    }

    import re as _re

    def _apply_named(text: str) -> str:
        """Sostituisce i segnaposto nominali {{chiave}} (case-insensitive)."""
        if "{{" not in text:
            return text
        def _sub(m):
            key = m.group(1).strip().lower()
            key = named_aliases.get(key, key)
            return str(data_values.get(key, m.group(0)))
        return _re.sub(r"\{\{\s*([\w]+)\s*\}\}", _sub, text)

    def replace_placeholders(text: str) -> str:
        """Replace ellipsis placeholders with employee data."""
        result = _apply_named(text)
        # Periodo di prova parametrico anche su template con "15 giorni" fisso.
        pp = data_values["periodo_prova"]
        if pp and pp != "______":
            result = _re.sub(r'(prova di|minimo di)\s*\d+\s*giorni',
                             rf'\1 {pp} giorni', result, flags=_re.IGNORECASE)
        if "…" not in result:
            return result

        # The template uses Unicode ellipsis character (…) repeated multiple times
        # We need to replace these patterns specifically
        
        # Pattern 1: "Lavoratore: ……………, nato a …………. il ……………………, residente in ………………………………… con codice fiscale ……………………………."
        if "Lavoratore:" in result and "…" in result:
            # Replace the entire line (niente mansione davanti al nome)
            result = f"Lavoratore: {data_values['nome_completo']}, nato a {data_values['luogo_nascita']} il {data_values['data_nascita']}, residente in {data_values['indirizzo']} con codice fiscale {data_values['codice_fiscale']}."
        
        # Pattern 2: "IL Sig. ……………………………. è assunto" - this line contains EVERYTHING
        elif "IL Sig." in result and "è assunto" in result and "…" in result:
            import re
            # Replace name
            result = re.sub(r'IL Sig\.\s*[…\.]+\s*è assunto', f"IL Sig. {data_values['nome_completo']} è assunto", result)
            # Replace mansioni
            result = re.sub(r'mansioni:\s*[…\.]+\s*inquadrato', f"mansioni: {data_values['mansione']} inquadrato", result, flags=re.IGNORECASE)
            # Replace livello
            result = re.sub(r'livello\s*[…\.]+\s*e con', f"livello {data_values['livello']} e con", result, flags=re.IGNORECASE)
            # Replace qualifica
            result = re.sub(r'qualifica\s*[…\.]+\s*del', f"qualifica {data_values['qualifica']} del", result, flags=re.IGNORECASE)
            # Replace date decorrenza (senza "al" se indeterminato)
            result = re.sub(r'decorre dal\s*[…\.]+\s*al\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
            result = re.sub(r'decorre dal\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
        
        # Pattern 3: "mansioni: ………………………… inquadrato"
        elif "mansioni:" in result and "…" in result:
            import re
            result = re.sub(r'mansioni:\s*[…\.]+\s*inquadrato', f"mansioni: {data_values['mansione']} inquadrato", result)
        
        # Pattern 3b: "delle seguenti mansioni:" followed by placeholders
        elif "seguenti mansioni:" in result.lower() and "…" in result:
            import re
            result = re.sub(r'seguenti mansioni:\s*[…\.]+', f"seguenti mansioni: {data_values['mansione']}", result, flags=re.IGNORECASE)
        
        # Pattern 4: "livello …….." or "livello …………"
        elif "livello" in result.lower() and "…" in result:
            import re
            result = re.sub(r'livello\s*[…\.]+', f"livello {data_values['livello']}", result, flags=re.IGNORECASE)
        
        # Pattern 5: "qualifica …………" 
        elif "qualifica" in result.lower() and "…" in result:
            import re
            result = re.sub(r'qualifica\s*[…\.]+', f"qualifica {data_values['qualifica']}", result, flags=re.IGNORECASE)
        
        # Pattern 6: "euro ………………… ora" (stipendio)
        elif "euro" in result.lower() and "ora" in result.lower() and "…" in result:
            import re
            result = re.sub(r'euro\s*[…\.]+\s*ora', f"euro {data_values['stipendio_orario']} ora", result, flags=re.IGNORECASE)
        
        # Pattern 7: "decorre dal ………… al …………"
        elif "decorre dal" in result.lower() and "…" in result:
            import re
            result = re.sub(r'decorre dal\s*[…\.]+\s*al\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
            result = re.sub(r'decorre dal\s*[…\.]+', _decorr, result, flags=re.IGNORECASE)
        
        # Generic fallback: replace any remaining ellipsis sequences
        elif "…" in result:
            import re
            # Replace sequences of 10+ ellipsis with longer values
            result = re.sub(r'[…]{10,}', data_values['nome_completo'], result)
            # Replace sequences of 6-9 ellipsis with medium values
            result = re.sub(r'[…]{6,9}', data_values['mansione'], result)
            # Replace sequences of 3-5 ellipsis with shorter values
            result = re.sub(r'[…]{3,5}', "______", result)
        
        return result
    
    # Process all paragraphs (ogni paragrafo: i segnaposto e le sostituzioni
    # parametriche come il periodo di prova possono comparire anche senza puntini).
    for para in doc.paragraphs:
        new_text = replace_placeholders(para.text)
        if new_text != para.text:
            if para.runs:
                first_run = para.runs[0]
                for run in para.runs[1:]:
                    run.text = ""
                first_run.text = new_text
            else:
                para.text = new_text

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text = replace_placeholders(para.text)
                    if new_text != para.text:
                        if para.runs:
                            first_run = para.runs[0]
                            for run in para.runs[1:]:
                                run.text = ""
                            first_run.text = new_text
                        else:
                            para.text = new_text
    
    # Save to temp file
    output_path = tempfile.mktemp(suffix=".docx")
    doc.save(output_path)
    
    return output_path


@router.get("/types")
@handle_errors
async def get_contract_types() -> List[Dict[str, str]]:
    """Get available contract types."""
    return CONTRACT_TYPES


# ---------------------------------------------------------------------------
# CCNL: livelli e retribuzioni. Servono all'anagrafica per non far digitare a
# mano gli importi e per accorgersi subito se una paga sta sotto il minimo.
# ---------------------------------------------------------------------------
@router.get("/ccnl")
@handle_errors
async def elenco_ccnl() -> List[Dict[str, Any]]:
    """CCNL disponibili. `tabelle_caricate` dice se si puo' gia' calcolare."""
    from app.hr.services.ccnl import lista_ccnl
    return lista_ccnl()


@router.get("/ccnl/{ccnl_id}/livello/{livello}")
@handle_errors
async def ccnl_retribuzione(ccnl_id: str, livello: str,
                            ore_settimanali: float = 40, scatti: int = 0) -> Dict[str, Any]:
    """Livello -> retribuzione mensile, giornaliera e oraria (part-time incluso)."""
    from app.hr.services.ccnl import retribuzione_per_livello, CCNLNonDisponibile
    try:
        return retribuzione_per_livello(livello, ccnl_id, ore_settimanali, scatti)
    except CCNLNonDisponibile as e:
        raise HTTPException(422, str(e))


@router.get("/ccnl/verifica-tranche")
@handle_errors
async def ccnl_verifica_tranche(ccnl: str = "turismo_pubblici_esercizi",
                                mesi: int = 12) -> Dict[str, Any]:
    """Confronta la paga base applicata (dalle buste) col tabellare, per livello.

    Risponde alla domanda "che tranche di rinnovo stiamo applicando": uno
    scarto comune a piu' livelli e' una tranche precedente, uno scarto isolato
    su un livello o una persona e' il caso da controllare col consulente.
    """
    from app.hr.services.verifica_tranche import verifica_tranche
    from app.hr.services.ccnl import CCNLNonDisponibile
    db = Database.get_db()
    try:
        return await verifica_tranche(db, ccnl, mesi)
    except CCNLNonDisponibile as e:
        raise HTTPException(422, str(e))


@router.post("/cedolini/importa-libro-unico")
@handle_errors
async def importa_libro_unico(file: UploadFile = File(...)) -> Dict[str, Any]:
    """Carica un Libro Unico multi-dipendente: lo divide per persona e registra
    i cedolini nuovi. Non duplica un mese gia' presente per lo stesso dipendente.

    E' il flusso normale con cui arrivano le buste dal consulente: un PDF unico
    con tutti i dipendenti del mese, pagine di presenze e di elementi
    retributivi in coppia per persona.
    """
    from app.hr.services.libro_unico_bundle import dividi_e_registra
    pdf_bytes = await file.read()
    db = Database.get_db()
    return await dividi_e_registra(db, pdf_bytes, file.filename or "")


@router.get("/acconti-tfr/{employee_id}")
@handle_errors
async def acconti_tfr_dipendente(employee_id: str) -> Dict[str, Any]:
    """Acconti sulla retribuzione e anticipi TFR di un dipendente, per anno."""
    from app.hr.services.acconti_tfr import riepilogo_dipendente
    db = Database.get_db()
    return await riepilogo_dipendente(db, employee_id)


@router.get("/acconti-tfr")
@handle_errors
async def acconti_tfr_azienda(anno: Optional[int] = None) -> List[Dict[str, Any]]:
    """Stessa cosa per tutti i dipendenti. `anno` per filtrare un anno solo."""
    from app.hr.services.acconti_tfr import riepilogo_azienda
    db = Database.get_db()
    return await riepilogo_azienda(db, anno)


@router.get("/profilo-retributivo/{employee_id}")
@handle_errors
async def profilo_retributivo(employee_id: str, ccnl: Optional[str] = None) -> Dict[str, Any]:
    """Cosa sappiamo davvero della paga di un dipendente.

    Unisce le sue buste (livello, lordo, netto), i bonifici realmente usciti dal
    conto e il minimo tabellare del CCNL, e segnala gli scostamenti. E' il dato
    con cui precompilare il contratto senza digitare importi a mano.
    """
    from app.hr.services.profilo_retributivo import profilo
    db = Database.get_db()
    emp = await db[Collections.EMPLOYEES].find_one({"id": employee_id}, {"_id": 0})
    if not emp:
        raise HTTPException(404, "Dipendente non trovato")
    return await profilo(db, emp, ccnl)


@router.post("/ccnl/suggerisci")
@handle_errors
async def ccnl_suggerisci(data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """Importo -> livello suggerito, con la classifica per scarto.

    Body: {"importo_mensile": 1600, "ccnl": "terziario", "ore_settimanali": 40}
    """
    from app.hr.services.ccnl import suggerisci_livello, CCNLNonDisponibile
    try:
        return suggerisci_livello(data.get("importo_mensile"),
                                  data.get("ccnl"),
                                  data.get("ore_settimanali"))
    except CCNLNonDisponibile as e:
        raise HTTPException(422, str(e))


@router.get("/templates")
@handle_errors
async def list_templates() -> List[Dict[str, Any]]:
    """Elenco template con disponibilità (MongoDB-first, poi disco)."""
    ensure_dirs()
    db = Database.get_db()
    in_mongo = {d["tipo"] async for d in db[COLL_TEMPLATES].find({}, {"_id": 0, "tipo": 1})}
    templates = []
    for ct in CONTRACT_TYPES:
        exists = (ct["id"] in in_mongo) or os.path.exists(os.path.join(TEMPLATES_DIR, ct["filename"]))
        templates.append({
            "id": ct["id"],
            "name": ct["name"],
            "filename": ct["filename"],
            "available": exists,
        })
    return templates


@router.post("/template/{contract_type}")
@handle_errors
async def upload_template(contract_type: str, file: UploadFile = File(...)) -> Dict[str, Any]:
    """Carica/sostituisce un template .docx (salvato su MongoDB, persistente)."""
    ct = next((c for c in CONTRACT_TYPES if c["id"] == contract_type), None)
    if not ct:
        raise HTTPException(400, f"Tipo contratto non valido: {contract_type}")
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "File vuoto")
    if len(raw) > 12 * 1024 * 1024:
        raise HTTPException(400, "File troppo grande (max 12MB)")
    if raw[:2] != b"PK":  # i .docx sono archivi ZIP (firma 'PK')
        raise HTTPException(400, "Il file non è un .docx valido")
    db = Database.get_db()
    await db[COLL_TEMPLATES].update_one(
        {"tipo": contract_type},
        {"$set": {
            "tipo": contract_type,
            "name": ct["name"],
            "filename": file.filename or ct["filename"],
            "file_data": base64.b64encode(raw).decode("utf-8"),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True)
    return {"ok": True, "tipo": contract_type, "name": ct["name"]}


# Documenti accessori generati automaticamente insieme al contratto.
ACCESSORI_AUTO = ["regolamento", "informativa_privacy", "informativa_152"]


async def _template_disponibile(db, ct_id: str) -> bool:
    """True se il template (MongoDB o disco) per quel tipo è caricato."""
    ct = next((c for c in CONTRACT_TYPES if c["id"] == ct_id), None)
    if not ct:
        return False
    doc = await db[COLL_TEMPLATES].find_one({"tipo": ct_id}, {"_id": 0, "tipo": 1})
    if doc:
        return True
    return os.path.exists(os.path.join(TEMPLATES_DIR, ct["filename"]))


async def _genera_doc(db, employee: Dict[str, Any], ct: Dict[str, str],
                      additional_data: Dict[str, Any]) -> Dict[str, Any]:
    """Compila un template per il dipendente, salva il record e lo ritorna.

    Cuore unico della generazione: usato dal contratto principale, dai documenti
    accessori e dalla generazione massiva.
    """
    ensure_dirs()
    template_path = await _resolve_template(ct)
    employee_data = {**employee, **(additional_data or {})}
    if employee_data.get("data_nascita"):
        try:
            dt = datetime.fromisoformat(str(employee_data["data_nascita"]).replace("Z", ""))
            employee_data["data_nascita"] = dt.strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            pass

    output_path = fill_contract_template(template_path, employee_data)
    nome_completo = employee_data.get("nome_completo") or \
        f"{employee_data.get('cognome','')} {employee_data.get('nome','')}".strip()
    safe_name = (nome_completo or "dipendente").replace(" ", "_")
    final_filename = f"{ct['id']}_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    final_path = os.path.join(CONTRACTS_DIR, final_filename)
    shutil.move(output_path, final_path)
    with open(final_path, "rb") as f:
        file_base64 = base64.b64encode(f.read()).decode("utf-8")

    mensile = compute_stipendio_mensile(
        employee_data.get("stipendio_orario") or employee_data.get("salary"),
        employee_data.get("ore_settimanali") or "40")

    record = {
        "id": str(uuid.uuid4()),
        "employee_id": employee.get("id"),
        "employee_name": nome_completo,
        "contract_type": ct["id"],
        "contract_name": ct["name"],
        "filename": final_filename,
        "filepath": final_path,
        "file_data": file_base64,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "stipendio_mensile": mensile,
        "iter_stato": "bozza",
        "additional_data": additional_data or {},
    }
    await db["employee_contracts"].insert_one(record.copy())
    return record


async def _genera_con_accessori(db, employee: Dict[str, Any], contract_type: str,
                                additional_data: Dict[str, Any]) -> Dict[str, Any]:
    """Genera il contratto principale e, se i template sono caricati, anche i
    documenti accessori (regolamento, privacy, informativa 152)."""
    ct = next((c for c in CONTRACT_TYPES if c["id"] == contract_type), None)
    if not ct:
        raise HTTPException(400, f"Tipo contratto non valido: {contract_type}")
    rec = await _genera_doc(db, employee, ct, additional_data)
    accessori, mancanti = [], []
    for acc_id in ACCESSORI_AUTO:
        acc = next((c for c in CONTRACT_TYPES if c["id"] == acc_id), None)
        if acc and await _template_disponibile(db, acc_id):
            a = await _genera_doc(db, employee, acc, additional_data)
            accessori.append(a["filename"])
        else:
            mancanti.append(acc_id)
    return {"contract": rec, "accessori": accessori, "accessori_mancanti": mancanti}


@router.post("/generate/{employee_id}")
@handle_errors
async def generate_contract(employee_id: str, data: Dict[str, Any] = Body(...)) -> Dict[str, Any]:
    """
    Generate a contract for an employee, insieme ai documenti accessori
    (regolamento, privacy, informativa 152) se i relativi template sono caricati.

    Request body:
    {
        "contract_type": "determinato",
        "additional_data": {"livello": "5", "stipendio_orario": "8.50", "qualifica": "Barista"}
    }
    """
    contract_type = data.get("contract_type") or data.get("contract_type_id")
    additional_data = data.get("additional_data", {})

    db = Database.get_db()
    employee = await db[Collections.EMPLOYEES].find_one({"id": employee_id}, {"_id": 0})
    if not employee:
        raise HTTPException(status_code=404, detail="Dipendente non trovato")

    try:
        res = await _genera_con_accessori(db, employee, contract_type, additional_data)
        rec = res["contract"]
        return {
            "success": True,
            "message": f"Contratto generato per {rec.get('employee_name')}",
            "stipendio_mensile": rec.get("stipendio_mensile"),
            "accessori": res["accessori"],
            "accessori_mancanti": res["accessori_mancanti"],
            "contract": {
                "id": rec["id"],
                "filename": rec["filename"],
                "download_url": f"/api/contracts/download/{rec['id']}"
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating contract: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Errore nella generazione del contratto. Riprova o contatta l'assistenza.")


def _deduci_tipo(employee: Dict[str, Any], ore_sett: Optional[float]) -> str:
    """Deduce il tipo di contratto da anagrafica + ore settimanali."""
    contratto = (employee.get("contratto") or "Indeterminato").lower()
    determinato = ("indeterminat" not in contratto and "determinat" in contratto) \
        or bool(employee.get("data_fine_contratto"))
    part = ore_sett is not None and ore_sett < 36
    if part and determinato:
        return "part_time_det"
    if part:
        return "part_time_ind"
    if determinato:
        return "determinato"
    return "indeterminato"


async def _dati_da_busta(db, employee: Dict[str, Any]) -> Dict[str, Any]:
    """Deduce livello, paga oraria e ore settimanali dall'ultima busta paga in
    archivio (collezione cedolini), abbinata per dipendente_id o codice fiscale."""
    cf = (employee.get("codice_fiscale") or "").strip()
    ors: List[Dict[str, Any]] = [{"dipendente_id": employee.get("id")}]
    if cf:
        ors += [{"codice_fiscale": cf}, {"codice_fiscale": cf.upper()}]
    try:
        ced = await db[Collections.PAYSLIPS].find_one(
            {"$or": ors}, {"_id": 0}, sort=[("anno", -1), ("mese", -1)])
    except Exception:
        ced = None
    out: Dict[str, Any] = {}
    if not ced:
        return out
    if ced.get("livello"):
        out["livello"] = str(ced["livello"])
    ore_mese = _to_float(ced.get("ore_lavorate"))
    lordo = _to_float(ced.get("lordo"))
    if ore_mese and ore_mese > 0:
        out["ore_settimanali"] = str(round(ore_mese / 4.333))
        if lordo:
            out["stipendio_orario"] = f"{lordo / ore_mese:.2f}".replace(".", ",")
    return out


@router.post("/genera-massivo")
@handle_errors
async def genera_massivo(data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Genera (e salva come bozza) i contratti per tutti i dipendenti in forza,
    deducendo tipo/ore/paga dall'ultima busta paga. Non invia nulla.

    Salta chi ha già un contratto generato (a meno di `force: true`) e chi non ha
    il template del tipo dedotto. NON invia né firma: sono bozze da rivedere.
    """
    db = Database.get_db()
    force = bool(data.get("force"))
    TIPI_CONTRATTO = {"indeterminato", "determinato", "part_time_det", "part_time_ind"}
    raw = await db[Collections.EMPLOYEES].find(
        {"merged_into": {"$exists": False}}, {"_id": 0}).to_list(1000)
    employees = [e for e in raw
                 if e.get("attivo", True) is not False
                 and (e.get("stato") or "attivo") not in ("cessato", "dismesso", "archiviato")]

    generati, saltati = [], []
    for emp in employees:
        nome = emp.get("nome_completo") or f"{emp.get('cognome','')} {emp.get('nome','')}".strip()
        try:
            if not force:
                gia = await db["employee_contracts"].find_one(
                    {"employee_id": emp.get("id"), "contract_type": {"$in": list(TIPI_CONTRATTO)}},
                    {"_id": 0, "id": 1})
                if gia:
                    saltati.append({"dipendente": nome, "motivo": "contratto già presente"})
                    continue
            busta = await _dati_da_busta(db, emp)
            ore = _to_float(busta.get("ore_settimanali"))
            tipo = _deduci_tipo(emp, ore)
            add = dict(busta)
            if emp.get("data_assunzione"):
                add.setdefault("data_inizio", str(emp["data_assunzione"])[:10])
            if emp.get("data_fine_contratto"):
                add.setdefault("data_fine", str(emp["data_fine_contratto"])[:10])
            if emp.get("ruolo"):
                add.setdefault("mansione", emp["ruolo"])
                add.setdefault("qualifica", emp["ruolo"])
            if not await _template_disponibile(db, tipo):
                saltati.append({"dipendente": nome, "motivo": f"template '{tipo}' non caricato"})
                continue
            res = await _genera_con_accessori(db, emp, tipo, add)
            generati.append({"dipendente": nome, "tipo": tipo,
                             "dati_da_busta": bool(busta),
                             "accessori_mancanti": res["accessori_mancanti"]})
        except HTTPException as e:
            saltati.append({"dipendente": nome, "motivo": str(e.detail)})
        except Exception as e:
            saltati.append({"dipendente": nome, "motivo": str(e)})

    return {"ok": True, "generati": len(generati), "saltati": len(saltati),
            "dettaglio": generati, "non_generati": saltati}


@router.get("/download/{contract_id}")
@handle_errors
async def download_contract(contract_id: str):
    """
    Download a generated contract.
    Architettura MongoDB-first: priorità a file_data da MongoDB.
    """
    import base64
    from fastapi.responses import Response
    
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    # Priorità: file_data da MongoDB (architettura MongoDB-first)
    file_data = contract.get("file_data")
    if file_data:
        content = base64.b64decode(file_data)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{contract.get("filename", "contratto.docx")}"'}
        )
    
    # Fallback per contratti legacy con solo filepath
    filepath = contract.get("filepath")
    if filepath and os.path.exists(filepath):
        return FileResponse(
            filepath,
            filename=contract.get("filename"),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    raise HTTPException(status_code=404, detail="File contratto non trovato")


@router.get("/employee/{employee_id}")
@handle_errors
async def get_employee_contracts(employee_id: str) -> List[Dict[str, Any]]:
    """Get all contracts for an employee."""
    db = Database.get_db()
    contracts = await db["employee_contracts"].find(
        {"employee_id": employee_id},
        {"_id": 0}
    ).sort("generated_at", -1).to_list(100)
    
    return contracts


@router.delete("/{contract_id}")
@handle_errors
async def delete_contract(contract_id: str) -> Dict[str, Any]:
    """
    Delete a generated contract.
    Architettura MongoDB-first: elimina dal database.
    """
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contratto non trovato")
    
    # Delete record from MongoDB (architettura MongoDB-first)
    await db["employee_contracts"].delete_one({"id": contract_id})
    
    # Cleanup opzionale: tenta eliminazione file locale se esiste (per retrocompatibilità)
    filepath = contract.get("filepath")
    if filepath:
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
        except Exception:
            pass  # Ignora errori filesystem, il dato importante è su MongoDB
    
    return {"success": True, "message": "Contratto eliminato"}


# ---------------------------------------------------------------------------
# Invio del contratto/regolamento per email al dipendente + presa visione
# ---------------------------------------------------------------------------
def _smtp_send(to_addr: str, subject: str, body: str, allegati: List[Dict[str, Any]]) -> None:
    """Invio email con allegati via SMTP. Credenziali SOLO da env Render."""
    host = os.getenv("SMTP_HOST", "smtp.gmail.com")
    port = int(os.getenv("SMTP_PORT", "465"))
    user = os.getenv("SMTP_EMAIL") or os.getenv("SMTP_USER")
    pwd = os.getenv("SMTP_PASSWORD")
    if not (user and pwd):
        raise HTTPException(503, "Email non configurata: imposta SMTP_EMAIL e SMTP_PASSWORD in env Render.")
    msg = EmailMessage()
    msg["From"] = user
    msg["To"] = to_addr
    msg["Subject"] = subject
    msg.set_content(body)
    for a in allegati:
        fn = a["filename"]
        subtype = "pdf" if fn.lower().endswith(".pdf") else \
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        msg.add_attachment(a["data"], maintype="application", subtype=subtype, filename=fn)
    if port == 465:
        with smtplib.SMTP_SSL(host, port, context=ssl.create_default_context()) as s:
            s.login(user, pwd)
            s.send_message(msg)
    else:
        with smtplib.SMTP(host, port) as s:
            s.starttls(context=ssl.create_default_context())
            s.login(user, pwd)
            s.send_message(msg)


# Documenti accessori che accompagnano SEMPRE il contratto (da sottoscrivere):
# informativa 152/1997, informativa privacy e regolamento interno.
DOC_ACCESSORI = ["informativa_152", "informativa_privacy", "regolamento"]


async def _raccogli_documenti(db, contract: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Contratto + ultimi documenti accessori generati per il dipendente
    (informativa 152, privacy, regolamento), nell'ordine di sottoscrizione.

    Ritorna [{filename, data(bytes), tipo}]. Solo i documenti effettivamente
    generati per quel dipendente vengono inclusi.
    """
    docs = [{"filename": contract.get("filename", "contratto.docx"),
             "data": base64.b64decode(contract["file_data"]),
             "tipo": contract.get("contract_type", "contratto")}]
    emp_id = contract.get("employee_id")
    for tipo in DOC_ACCESSORI:
        d = await db["employee_contracts"].find_one(
            {"employee_id": emp_id, "contract_type": tipo},
            {"_id": 0, "file_data": 1, "filename": 1}, sort=[("generated_at", -1)])
        if d and d.get("file_data"):
            docs.append({"filename": d.get("filename", f"{tipo}.docx"),
                         "data": base64.b64decode(d["file_data"]), "tipo": tipo})
    return docs


@router.post("/send/{contract_id}")
@handle_errors
async def send_contract(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Invia al dipendente via email il contratto INSIEME a regolamento, privacy
    e informativa 152 (quelli generati per lui), da sottoscrivere."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    emp = await db[Collections.EMPLOYEES].find_one(
        {"id": contract.get("employee_id")},
        {"_id": 0, "nome": 1, "cognome": 1, "email": 1}) or {}
    to_addr = (data.get("email") or emp.get("email") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Email del dipendente mancante: inseriscila in anagrafica.")
    if not contract.get("file_data"):
        raise HTTPException(404, "File del contratto non disponibile")

    documenti = await _raccogli_documenti(db, contract)
    allegati = [{"filename": d["filename"], "data": d["data"]} for d in documenti]
    mancanti = [t for t in DOC_ACCESSORI if t not in {d["tipo"] for d in documenti}]

    nome = f"{emp.get('nome','')} {emp.get('cognome','')}".strip() or "Gentile collaboratore"
    elenco = ", ".join(d["filename"] for d in documenti)
    corpo = (
        f"Gentile {nome},\n\n"
        f"in allegato trova i documenti di assunzione da sottoscrivere: {elenco}.\n"
        f"La preghiamo di firmarli per accettazione e di restituirli a Ceraldi Group, che "
        f"provvederà alla controfirma e all'invio della copia definitiva.\n\n"
        f"Ceraldi Group S.r.l."
    )
    import asyncio
    await asyncio.to_thread(_smtp_send, to_addr,
                            f"Documenti di assunzione da firmare — {contract.get('contract_name','')}", corpo, allegati)
    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {"inviato_il": datetime.now(timezone.utc).isoformat(), "inviato_a": to_addr,
                  "iter_stato": "inviata"}})
    return {"ok": True, "inviato_a": to_addr, "allegati": len(allegati),
            "documenti": [d["filename"] for d in documenti], "accessori_mancanti": mancanti}


# ---------------------------------------------------------------------------
# Iter di sottoscrizione (manuale): bozza -> inviata -> firmato_dipendente ->
# definitivo (controfirma Ceraldi + invio + archiviazione nel fascicolo).
# ---------------------------------------------------------------------------
ITER_STATI = ["bozza", "inviata", "firmato_dipendente", "definitivo"]


@router.post("/carica-firmato/{contract_id}")
@handle_errors
async def carica_firmato(contract_id: str, file: UploadFile = File(...)) -> Dict[str, Any]:
    """Carica il contratto FIRMATO dal dipendente e restituito a Ceraldi (PDF).
    Porta lo stato a 'firmato_dipendente'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0, "id": 1})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    raw = await file.read()
    if not raw:
        raise HTTPException(400, "File vuoto")
    if len(raw) > 20 * 1024 * 1024:
        raise HTTPException(400, "File troppo grande (max 20MB)")
    if raw[:4] != b"%PDF":
        raise HTTPException(400, "Il contratto firmato deve essere un PDF")
    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "pdf_firmato_dipendente": base64.b64encode(raw).decode("ascii"),
            "firmato_filename": file.filename or "contratto_firmato.pdf",
            "iter_stato": "firmato_dipendente",
            "firmato_dipendente_il": datetime.now(timezone.utc).isoformat(),
        }})
    return {"ok": True, "stato": "firmato_dipendente"}


@router.post("/finalizza/{contract_id}")
@handle_errors
async def finalizza_contratto(contract_id: str, file: Optional[UploadFile] = File(default=None),
                              pec: str = "", email: str = "") -> Dict[str, Any]:
    """Controfirma Ceraldi + invio definitivo + archiviazione nel fascicolo.

    Se `file` è presente è il PDF controfirmato da Ceraldi (definitivo); altrimenti
    si usa il PDF firmato dal dipendente. Invia la copia definitiva al dipendente
    (email e, se indicata, PEC) e la salva nel fascicolo (contratti_dipendenti).
    Porta lo stato a 'definitivo'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")

    if file is not None:
        definitivo = await file.read()
        if definitivo[:4] != b"%PDF":
            raise HTTPException(400, "Il file controfirmato deve essere un PDF")
        def_name = file.filename or "contratto_definitivo.pdf"
    elif contract.get("pdf_firmato_dipendente"):
        definitivo = base64.b64decode(contract["pdf_firmato_dipendente"])
        def_name = contract.get("firmato_filename", "contratto_definitivo.pdf")
    else:
        raise HTTPException(400, "Manca il contratto firmato: caricalo prima, o allega il PDF controfirmato.")
    if not def_name.lower().endswith(".pdf"):
        def_name += ".pdf"

    emp = await db[Collections.EMPLOYEES].find_one(
        {"id": contract.get("employee_id")}, {"_id": 0, "nome": 1, "cognome": 1, "email": 1}) or {}
    to_addr = (email or emp.get("email") or "").strip()

    def_b64 = base64.b64encode(definitivo).decode("ascii")
    now = datetime.now(timezone.utc).isoformat()

    # Invio definitivo al dipendente (email + PEC opzionale)
    inviato = []
    nome = f"{emp.get('nome','')} {emp.get('cognome','')}".strip() or "Gentile collaboratore"
    corpo = (f"Gentile {nome},\n\nin allegato la copia definitiva del contratto "
             f"({contract.get('contract_name','')}), controfirmata da Ceraldi Group.\n\n"
             f"Ceraldi Group S.r.l.")
    if to_addr:
        import asyncio
        await asyncio.to_thread(_smtp_send, to_addr,
                                f"Contratto definitivo — {contract.get('contract_name','')}",
                                corpo, [{"filename": def_name, "data": definitivo}])
        inviato.append(to_addr)
    if pec:
        client = get_client()
        if client.configured:
            try:
                await client.send_pec(to_addr=pec,
                                      subject=f"Contratto definitivo — {contract.get('contract_name','')}",
                                      body=corpo, attachments=[{"filename": def_name, "content": definitivo}])
                inviato.append(f"PEC:{pec}")
            except OpenAPIError as e:
                logger.warning(f"PEC finalizza: {e}")

    # Archiviazione nel fascicolo del dipendente (collezione contratti_dipendenti)
    add = contract.get("additional_data", {}) or {}
    fasc = {
        "id": str(uuid.uuid4()),
        "dipendente_id": contract.get("employee_id"),
        "tipo_contratto": contract.get("contract_type"),
        "nome": contract.get("contract_name"),
        "data_inizio": add.get("data_inizio"),
        "data_fine": add.get("data_fine") or None,
        "stato": "attivo",
        "firmato": True,
        "filename": def_name,
        "file_data": def_b64,
        "contract_ref": contract_id,
        "archiviato_il": now,
    }
    await db["contratti_dipendenti"].insert_one(fasc.copy())

    # A→B: il contratto definitivo entra ANCHE nei Documenti del dipendente (cartella),
    # così anagrafica, fascicolo e documenti restano collegati. Dedup per hash.
    try:
        import hashlib
        ctype0 = contract.get("contract_type", "")
        h = hashlib.sha256(definitivo).hexdigest()
        if not await db["documenti_cloud"].find_one({"hash": h}):
            cat = "RIDUZIONE_ORARIO" if "riduzione" in ctype0 else "CONTRATTO"
            await db["documenti_cloud"].insert_one({
                "id": str(uuid.uuid4()),
                "dipendente_id": contract.get("employee_id"),
                "dipendente_nome": nome,
                "titolo": contract.get("contract_name") or def_name,
                "filename": def_name, "tipo": cat, "categoria": cat, "hash": h,
                "file_data": def_b64, "assegnato": True,
                "origine": "contratto_finalizzato", "data_caricamento": now,
            })
    except Exception:
        logger.warning("documenti_cloud da contratto finalizzato: skip")

    # A→B: il contratto definitivo aggiorna l'anagrafica del dipendente.
    add_dati = contract.get("additional_data", {}) or {}
    ctype = contract.get("contract_type", "")
    emp_update: Dict[str, Any] = {}
    for campo, valore in {
        "livello": add_dati.get("livello"),
        "mansione": add_dati.get("mansione"),
        "ruolo": add_dati.get("mansione"),  # DipendenteCloud usa 'ruolo'
        "qualifica": add_dati.get("qualifica"),
        "stipendio_orario": add_dati.get("stipendio_orario"),
        "ore_settimanali": add_dati.get("ore_settimanali"),
        "data_assunzione": add_dati.get("data_inizio"),
        "data_fine_contratto": add_dati.get("data_fine"),
        "contratto": "Determinato" if ("determinato" in ctype and "indeterminato" not in ctype) else "Indeterminato",
    }.items():
        if valore not in (None, ""):
            emp_update[campo] = valore
    if emp_update:
        await db[Collections.EMPLOYEES].update_one(
            {"id": contract.get("employee_id")}, {"$set": emp_update})

    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "pdf_definitivo": def_b64,
            "definitivo_filename": def_name,
            "iter_stato": "definitivo",
            "definitivo_il": now,
            "definitivo_inviato_a": inviato,
            "fascicolo_id": fasc["id"],
        }})
    return {"ok": True, "stato": "definitivo", "inviato_a": inviato, "archiviato": True}


@router.get("/pdf/{contract_id}/{versione}")
@handle_errors
async def download_pdf_versione(contract_id: str, versione: str):
    """Scarica il PDF firmato dal dipendente o quello definitivo controfirmato."""
    campo = {"firmato": "pdf_firmato_dipendente", "definitivo": "pdf_definitivo"}.get(versione)
    if not campo:
        raise HTTPException(400, "Versione non valida (firmato|definitivo)")
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract or not contract.get(campo):
        raise HTTPException(404, "PDF non disponibile")
    name = contract.get("definitivo_filename" if versione == "definitivo" else "firmato_filename",
                        f"contratto_{versione}.pdf")
    return Response(content=base64.b64decode(contract[campo]), media_type="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{name}"'})


# ---------------------------------------------------------------------------
# Firma digitale via OpenAPI: marca temporale -> eSignature (FES+OTP) -> PEC
# Stato nel fascicolo: inviato -> firmato -> accettato.
# ---------------------------------------------------------------------------
def _docx_bytes_to_pdf(docx_bytes: bytes, filename: str = "contratto.docx") -> bytes:
    """Converte un .docx in PDF tramite il servizio unico (ConvertAPI in
    produzione, LibreOffice in locale). Vedi services/docx_converter.py."""
    try:
        return docx_to_pdf(docx_bytes, filename)
    except DocxConversionError as e:
        # 503: configurazione/servizio mancante (azionabile dal titolare).
        raise HTTPException(503, str(e))


async def _get_contract_pdf(contract: Dict[str, Any]) -> bytes:
    """Ritorna il PDF del contratto: usa quello già marcato/firmato se presente,
    altrimenti converte il .docx generato."""
    if contract.get("pdf_data"):
        return base64.b64decode(contract["pdf_data"])
    if not contract.get("file_data"):
        raise HTTPException(404, "File del contratto non disponibile")
    return _docx_bytes_to_pdf(base64.b64decode(contract["file_data"]))


def _bundle_to_pdf(documenti: List[Dict[str, Any]]) -> bytes:
    """Converte ogni .docx (contratto + accessori) in PDF e li unisce in un unico
    PDF: così la firma per accettazione copre tutti i documenti in una sola volta."""
    from PyPDF2 import PdfMerger
    merger = PdfMerger()
    for d in documenti:
        merger.append(io.BytesIO(_docx_bytes_to_pdf(d["data"], d.get("filename", "documento.docx"))))
    out = io.BytesIO()
    merger.write(out)
    merger.close()
    return out.getvalue()


@router.post("/sign/{contract_id}")
@handle_errors
async def avvia_firma(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Avvia il flusso di firma: converte in PDF, applica marca temporale e crea
    la richiesta di eSignature (FES con OTP) verso il dipendente."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    emp = await db[Collections.EMPLOYEES].find_one(
        {"id": contract.get("employee_id")},
        {"_id": 0, "nome": 1, "cognome": 1, "email": 1, "telefono": 1, "cellulare": 1}) or {}
    to_addr = (data.get("email") or emp.get("email") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Email del dipendente mancante: inseriscila in anagrafica.")
    phone = (data.get("phone") or emp.get("cellulare") or emp.get("telefono") or "").strip()
    nome = f"{emp.get('nome','')} {emp.get('cognome','')}".strip() or "Dipendente"

    client = get_client()
    if not client.configured:
        raise HTTPException(503, "OpenAPI non configurato: imposta OPENAPI_CLIENT_ID e "
                                 "OPENAPI_CLIENT_SECRET nelle env di Render.")
    import asyncio
    documenti = await _raccogli_documenti(db, contract)
    fname = (contract.get("filename", "documenti").rsplit(".", 1)[0]) + "_assunzione.pdf"
    try:
        # Unico PDF con contratto + regolamento + privacy + informativa 152.
        pdf_bytes = await asyncio.to_thread(_bundle_to_pdf, documenti)
        # 1) Marca temporale (data certa)
        ts = await client.apply_timestamp(pdf_bytes, filename=fname)
        # 2) eSignature FES con OTP verso il dipendente (firma unica su tutti i documenti)
        sig = await client.create_signature_request(
            pdf_bytes, signer_name=nome, signer_email=to_addr, signer_phone=phone,
            title=f"Documenti di assunzione — {contract.get('contract_name','')}",
            filename=fname)
    except OpenAPIConfigError as e:
        raise HTTPException(503, str(e))
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    req_id = sig.get("id") or sig.get("request_id")
    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "pdf_data": base64.b64encode(pdf_bytes).decode("ascii"),
            "firma_stato": "inviato",
            "firma_request_id": req_id,
            "marca_temporale": ts,
            "firma_inviata_a": to_addr,
            "firma_inviata_il": datetime.now(timezone.utc).isoformat(),
            "firma_documenti": [d["filename"] for d in documenti],
        }})
    return {"ok": True, "stato": "inviato", "request_id": req_id, "firmatario": to_addr,
            "documenti": [d["filename"] for d in documenti]}


@router.get("/sign/{contract_id}/status")
@handle_errors
async def stato_firma(contract_id: str) -> Dict[str, Any]:
    """Interroga OpenAPI sullo stato della firma e aggiorna il fascicolo.
    Se firmato, salva il PDF firmato e passa lo stato a 'firmato'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    req_id = contract.get("firma_request_id")
    if not req_id:
        return {"ok": True, "stato": contract.get("firma_stato") or "non_avviato"}

    client = get_client()
    try:
        res = await client.get_signature_status(req_id)
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    stato_raw = str(res.get("status") or "").lower()
    updates: Dict[str, Any] = {"firma_check_il": datetime.now(timezone.utc).isoformat()}
    firmato = stato_raw in ("signed", "completed", "firmato")
    if firmato:
        updates["firma_stato"] = "firmato"
        signed = res.get("signed_document") or res.get("signed_pdf") or {}
        content = signed.get("content") if isinstance(signed, dict) else signed
        if content:
            updates["pdf_data"] = content  # base64 del PDF firmato
            updates["firmato_il"] = datetime.now(timezone.utc).isoformat()
    await db["employee_contracts"].update_one({"id": contract_id}, {"$set": updates})
    return {"ok": True, "stato": updates.get("firma_stato", contract.get("firma_stato")), "provider": stato_raw}


@router.post("/pec/{contract_id}")
@handle_errors
async def invia_pec(contract_id: str, data: Dict[str, Any] = Body(default={})) -> Dict[str, Any]:
    """Invia via PEC il contratto (firmato se disponibile) con ricevuta a data certa.
    Porta lo stato a 'accettato'."""
    db = Database.get_db()
    contract = await db["employee_contracts"].find_one({"id": contract_id}, {"_id": 0})
    if not contract:
        raise HTTPException(404, "Contratto non trovato")
    to_addr = (data.get("pec") or "").strip()
    if not to_addr:
        raise HTTPException(400, "Indirizzo PEC destinatario mancante.")
    pdf_bytes = await _get_contract_pdf(contract)
    pdf_name = (contract.get("filename", "contratto.docx").rsplit(".", 1)[0]) + ".pdf"

    client = get_client()
    if not client.configured:
        raise HTTPException(503, "OpenAPI non configurato (OPENAPI_CLIENT_ID/SECRET in env Render).")
    try:
        res = await client.send_pec(
            to_addr=to_addr,
            subject=f"Contratto di assunzione — {contract.get('contract_name','')}",
            body="In allegato il contratto di assunzione con marca temporale e firma per accettazione.",
            attachments=[{"filename": pdf_name, "content": pdf_bytes}])
    except OpenAPIConfigError as e:
        raise HTTPException(503, str(e))
    except OpenAPIError as e:
        raise HTTPException(502, f"OpenAPI: {e}")

    await db["employee_contracts"].update_one(
        {"id": contract_id},
        {"$set": {
            "firma_stato": "accettato",
            "pec_inviata_a": to_addr,
            "pec_inviata_il": datetime.now(timezone.utc).isoformat(),
            "pec_messaggio": res,
        }})
    return {"ok": True, "stato": "accettato", "pec": to_addr}
